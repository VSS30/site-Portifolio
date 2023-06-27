import docx

def gera_word():
    # Cria um novo documento do Word
    doc = docx.Document()
    
    section = doc.sections[0]
   
    # Adicionar bordas às tabelas
    border = docx.enum.table.WD_BORDER.SINGLE
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'Table Normal'
                    for run in paragraph.runs:
                        run.font.size = docx.shared.Pt(12)
                        run.font.name = 'Arial'
                        cell._element.tcPr.append(docx.oxml.shared.OxmlElement('w:tcBorders'))
                        cell._element.tcPr.tcBorders.append(
                            docx.oxml.shared.OxmlElement('w:top', {'w:val': border})
                        )
                        cell._element.tcPr.tcBorders.append(
                            docx.oxml.shared.OxmlElement('w:left', {'w:val': border})
                        )
                        cell._element.tcPr.tcBorders.append(
                            docx.oxml.shared.OxmlElement('w:bottom', {'w:val': border})
                        )
                        cell._element.tcPr.tcBorders.append(
                            docx.oxml.shared.OxmlElement('w:right', {'w:val': border})
                        )

    
    
    section.page_background = 'C:\\Users\\Vinicius\\Ecomerce_API_ICO\\static\\LOGO.png'

    # Adiciona um título
    titulo = doc.add_paragraph(style='Heading 1')
    titulo.add_run('Ordem de Serviço')

    # Adiciona informações do cliente
    info_cliente = doc.add_paragraph(style='Heading 2')
    info_cliente.add_run('Informações do Cliente')

  
  
    doc.add_paragraph('Nome do Cliente:: João')
    doc.add_paragraph('Endereço: Rua A, 123')
    doc.add_paragraph('Telefone: (11) 1234-5678')
    doc.add_paragraph('E-mail: Vini.baxista@outlook.com.br')
   

    # Adiciona informações do orçamento
    info_orcamento = doc.add_paragraph(style='Heading 2')
    info_orcamento.add_run('Informações do Orçamento')

   

    # Preenche a tabela com informações do orçamento
    servico_table = doc.add_table(rows=4, cols=2)
    servico_table.style = 'Table Grid'
    servico_table.cell(0, 0).text = 'Modelo do Celular:'
    servico_table.cell(1, 0).text = 'Problema:'
    servico_table.cell(2, 0).text = 'Data de Entrega:'
    servico_table.cell(3, 0).text = 'Valor:'
    servico_table.cell(0, 1).text = 'iPhone X'
    servico_table.cell(1, 1).text = 'Tela Quebrada'
    servico_table.cell(2, 1).text = '01/04/2023'
    servico_table.cell(3, 1).text = 'R$ 500,00'

    # Adiciona uma imagem
    # doc.add_picture('C:\\Users\\Vinicius\\Ecomerce_API_ICO\\static\\LOGO.png', width=docx.shared.Inches(2))

    # Salva o documento em um arquivo .docx
    doc.save('orcamento.docx')


gera_word()